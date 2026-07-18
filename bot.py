import os
<<<<<<< HEAD
import json
=======
>>>>>>> af8b08ecffec43e5dc1c5577f5759f7b6403b00a
import pickle
from pathlib import Path

from dotenv import load_dotenv
from langchain_core.prompts import ChatPromptTemplate
from langchain.memory import ConversationBufferMemory
from langchain_groq import ChatGroq
from langchain_astradb import AstraDBVectorStore
from langchain_huggingface import HuggingFaceEndpointEmbeddings
from langchain_community.retrievers import BM25Retriever
from langchain.retrievers import EnsembleRetriever
<<<<<<< HEAD
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
import pandas as pd
import docx
=======
>>>>>>> af8b08ecffec43e5dc1c5577f5759f7b6403b00a

load_dotenv()  # reads .env from the current working directory

COLLECTION_NAME = "sales_products"          # AstraDB collection created by ingest.py
EMBEDDING_MODEL = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
BM25_CACHE = Path("bm25_corpus.pkl")        # local chunk cache created by ingest.py
<<<<<<< HEAD
DATA_DIR = Path("Data")
METADATA_FILE = Path("ingested_metadata.json")
=======
>>>>>>> af8b08ecffec43e5dc1c5577f5759f7b6403b00a


PROMPT_TEMPLATE = ChatPromptTemplate.from_messages([
    (
        "system",
        """
You are a professional AI Sales Assistant developed by Keshav Gupta.

Your primary goal is to help customers make informed purchasing decisions using ONLY the provided product information.

Guidelines:
1. Always communicate in clear, professional, and friendly English.
2. Answer only using the provided Product Information.
3. Never make up product specifications, prices, availability, or features.
4. If the required information is not available in the Product Information, politely respond:
   "I'm sorry, but I couldn't find that information in the available product data."
5. Recommend the most suitable product based on the customer's:
   - Budget
   - Requirements
   - Use case
   - Preferences
6. When multiple products match the customer's needs:
   - Compare them in detail.
   - Explain the advantages and disadvantages of each.
   - Recommend the best option and explain why.
7. Clearly explain:
   - Key features
   - Benefits
   - Technical specifications
   - Suitable use cases
8. If the customer is unsure what to buy, ask relevant follow-up questions before recommending a product.
9. If a comparison is requested, present the comparison in a clean table whenever possible.
10. Be concise, accurate, and customer-focused.
11. Never mention internal documents, embeddings, vector databases, or retrieval systems.
12. Do not generate false information or hallucinate answers.

Product Information:
{context}
"""
    ),
    ("human", "{input}")
])

# Everything below is set up once, the first time getAnswer() is called —
# not at import time — so `from bot import getAnswer` never crashes on its own.
_state = {"initialized": False}


<<<<<<< HEAD
def load_file(file_path: Path) -> list[Document]:
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        try:
            loader = PyPDFLoader(str(file_path))
            return loader.load()
        except Exception as e:
            print(f"Error loading PDF {file_path}: {e}")
            return []
    elif ext == ".docx":
        try:
            doc = docx.Document(file_path)
            fullText = [para.text for para in doc.paragraphs if para.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        fullText.append(row_text)
            text = "\n".join(fullText)
            return [Document(page_content=text, metadata={"source": str(file_path)})]
        except Exception as e:
            print(f"Error loading docx {file_path}: {e}")
            return []
    elif ext in [".xlsx", ".xls"]:
        try:
            xls = pd.ExcelFile(file_path)
            docs = []
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                rows_text = []
                for idx, row in df.iterrows():
                    row_str = ", ".join(f"{col}: {val}" for col, val in row.items() if pd.notna(val))
                    if row_str.strip():
                        rows_text.append(f"Row {idx+1}: {row_str}")
                sheet_text = f"File: {file_path.name}, Sheet: {sheet_name}\n" + "\n".join(rows_text)
                if sheet_text.strip():
                    docs.append(Document(page_content=sheet_text, metadata={"source": str(file_path), "sheet": sheet_name, "row": idx}))
            return docs
        except Exception as e:
            print(f"Error loading excel {file_path}: {e}")
            return []
    elif ext in [".txt", ".csv", ".json"]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            return [Document(page_content=text, metadata={"source": str(file_path)})]
        except Exception as e:
            print(f"Error loading text file {file_path}: {e}")
            return []
    else:
        return []


def sync_data_folder(vstore):
    """Scan the Data/ directory, find new/modified/deleted files, load/chunk/ingest them, 
    and delete old/stale chunks from AstraDB and BM25.
    """
    if not DATA_DIR.exists():
        DATA_DIR.mkdir(parents=True, exist_ok=True)
        return

    # Load metadata of ingested files
    metadata = {}
    if METADATA_FILE.exists():
        try:
            with open(METADATA_FILE, "r") as f:
                metadata = json.load(f)
        except Exception as e:
            print(f"Error loading metadata file: {e}")
            metadata = {}

    # Load existing BM25 corpus if it exists
    all_chunks = []
    if BM25_CACHE.exists():
        try:
            with open(BM25_CACHE, "rb") as f:
                all_chunks = pickle.load(f)
        except Exception as e:
            print(f"Error loading BM25 cache: {e}")
            all_chunks = []

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    updated = False

    # Get list of current files in the Data/ directory
    current_files = {}
    for file_path in DATA_DIR.iterdir():
        if file_path.is_dir() or file_path.name.startswith("."):
            continue
        file_stat = file_path.stat()
        current_files[file_path.name] = {
            "path": file_path,
            "mtime": file_stat.st_mtime,
            "size": file_stat.st_size
        }

    # 1. Handle Deleted Files: Files in metadata but not in current_files
    deleted_files = [f_name for f_name in metadata if f_name not in current_files]
    for file_name in deleted_files:
        print(f"File deleted: {file_name}. Deleting from databases...")
        file_info = metadata[file_name]
        
        # Delete from AstraDB
        doc_ids = file_info.get("doc_ids", [])
        if doc_ids:
            try:
                vstore.delete(ids=doc_ids)
                print(f"Deleted {len(doc_ids)} chunks from AstraDB for {file_name}")
            except Exception as e:
                print(f"Error deleting from AstraDB for {file_name}: {e}")
        
        # Delete from BM25 corpus
        all_chunks = [c for c in all_chunks if not (c.metadata.get("source") and Path(c.metadata.get("source")).name == file_name)]
        del metadata[file_name]
        updated = True

    # 2. Handle New and Modified Files
    for file_name, file_info in current_files.items():
        file_path = file_info["path"]
        mtime = file_info["mtime"]
        size = file_info["size"]

        needs_ingestion = False
        if file_name not in metadata:
            needs_ingestion = True
        else:
            saved_info = metadata[file_name]
            if saved_info.get("mtime") != mtime or saved_info.get("size") != size:
                needs_ingestion = True
                # If modified, delete the old chunks first
                print(f"File modified: {file_name}. Deleting old data...")
                old_doc_ids = saved_info.get("doc_ids", [])
                if old_doc_ids:
                    try:
                        vstore.delete(ids=old_doc_ids)
                        print(f"Deleted {len(old_doc_ids)} old chunks from AstraDB for {file_name}")
                    except Exception as e:
                        print(f"Error deleting old chunks from AstraDB for {file_name}: {e}")
                
                # Delete old chunks from BM25 corpus
                all_chunks = [c for c in all_chunks if not (c.metadata.get("source") and Path(c.metadata.get("source")).name == file_name)]

        if needs_ingestion:
            print(f"Ingesting file: {file_name}")
            docs = load_file(file_path)
            if not docs:
                continue

            chunks = text_splitter.split_documents(docs)
            added_ids = []
            if chunks:
                # Normalize chunk metadata source
                for chunk in chunks:
                    chunk.metadata["source"] = str(file_path)

                # Add to AstraDB Vector Store and get IDs
                try:
                    added_ids = vstore.add_documents(chunks)
                    print(f"Added {len(chunks)} chunks to AstraDB for {file_name}")
                except Exception as e:
                    print(f"Error adding to AstraDB for {file_name}: {e}")
                    added_ids = []

                # Add to BM25 list
                all_chunks.extend(chunks)
                updated = True

            # Save file metadata including document IDs
            metadata[file_name] = {
                "mtime": mtime,
                "size": size,
                "doc_ids": added_ids
            }

    # Save BM25 corpus and metadata files if anything changed, or if files don't exist
    if updated or not BM25_CACHE.exists():
        if not all_chunks:
            # Add a placeholder chunk if the corpus is completely empty to prevent retriever crash
            all_chunks = [Document(page_content="No product information available yet.", metadata={"source": "none"})]

        try:
            with open(BM25_CACHE, "wb") as f:
                pickle.dump(all_chunks, f)
            print(f"Saved {len(all_chunks)} chunks to {BM25_CACHE}")
        except Exception as e:
            print(f"Error saving BM25 cache: {e}")

        try:
            with open(METADATA_FILE, "w") as f:
                json.dump(metadata, f, indent=2)
            print(f"Saved metadata to {METADATA_FILE}")
        except Exception as e:
            print(f"Error saving metadata file: {e}")


=======
>>>>>>> af8b08ecffec43e5dc1c5577f5759f7b6403b00a
def _initialize():
    """One-time setup: LLM, memory, and the hybrid (AstraDB + BM25) retriever."""
    if _state["initialized"]:
        return

    groq_api_key = os.getenv("GROQ_API")
    astra_db_url = os.getenv("CHATBOT_DB_URL")
    astra_db_token = os.getenv("CHATBOT_DB_TOKEN")
    hf_token = os.getenv("HUGGINGFACEHUB_API_TOKEN")

    # 🔹 LLM
    _state["llm"] = ChatGroq(
        groq_api_key=groq_api_key,
        model="llama-3.3-70b-versatile",
        temperature=0,
    )

    # 🔹 Memory
    _state["memory"] = ConversationBufferMemory(return_messages=True)

    # 🔹 AstraDB — dense/semantic search (hosted embeddings, no local torch needed)
    embeddings = HuggingFaceEndpointEmbeddings(
        model=EMBEDDING_MODEL,
        huggingfacehub_api_token=hf_token,
    )
    vstore = AstraDBVectorStore(
        embedding=embeddings,
        collection_name=COLLECTION_NAME,
        api_endpoint=astra_db_url,
        token=astra_db_token,
    )
<<<<<<< HEAD

    # Sync Data/ folder with databases (AstraDB & BM25)
    sync_data_folder(vstore)

    dense_retriever = vstore.as_retriever(search_kwargs={"k": 5})

    # 🔹 BM25 — sparse/keyword search, built from the same chunks
    if not BM25_CACHE.exists():
        all_chunks = [Document(page_content="No product information available yet.", metadata={"source": "none"})]
        with open(BM25_CACHE, "wb") as f:
            pickle.dump(all_chunks, f)
    else:
        with open(BM25_CACHE, "rb") as f:
            all_chunks = pickle.load(f)
=======
    dense_retriever = vstore.as_retriever(search_kwargs={"k": 5})

    # 🔹 BM25 — sparse/keyword search, built from the same chunks ingest.py embedded
    if not BM25_CACHE.exists():
        raise FileNotFoundError(
            "bm25_corpus.pkl not found. Run `python ingest.py` first to build the knowledge base."
        )
    with open(BM25_CACHE, "rb") as f:
        all_chunks = pickle.load(f)
>>>>>>> af8b08ecffec43e5dc1c5577f5759f7b6403b00a

    bm25_retriever = BM25Retriever.from_documents(all_chunks)
    bm25_retriever.k = 5

    # 🔹 Hybrid retriever = dense (meaning-based) + sparse (exact keyword) combined
    _state["retriever"] = EnsembleRetriever(
        retrievers=[dense_retriever, bm25_retriever],
        weights=[0.6, 0.4],
    )

    _state["initialized"] = True


def _format_docs(docs) -> str:
    if not docs:
        return "Koi matching information nahi mili."
    return "\n\n".join(f"[{i + 1}] {d.page_content}" for i, d in enumerate(docs))


# 🔹 The one function everything else imports and calls
def getAnswer(question: str) -> str:
    _initialize()

    llm = _state["llm"]
    memory = _state["memory"]
    retriever = _state["retriever"]

    # Retrieve previous chat history
    history = memory.buffer_as_messages
    formatted_history = "\n".join([f"{msg.type.capitalize()}: {msg.content}" for msg in history])

    # Hybrid RAG retrieval: combine AstraDB semantic search + BM25 keyword search
    retrieved_docs = retriever.invoke(question)
    context = _format_docs(retrieved_docs)

    # Build prompt, keeping prior chat history inline
    human_input = f"{formatted_history}\nUser: {question}" if formatted_history else question

    messages = PROMPT_TEMPLATE.format_messages(
        context=context,
        input=human_input,
    )

    response = ""
    # Stream response from Groq
    for chunk in llm.stream(messages):
        response += chunk.content

    # Save to memory
    memory.save_context({"input": question}, {"output": response})
    return response


<<<<<<< HEAD
def getAnswerStream(question: str):
    _initialize()

    llm = _state["llm"]
    memory = _state["memory"]
    retriever = _state["retriever"]

    # Retrieve previous chat history
    history = memory.buffer_as_messages
    formatted_history = "\n".join([f"{msg.type.capitalize()}: {msg.content}" for msg in history])

    # Hybrid RAG retrieval: combine AstraDB semantic search + BM25 keyword search
    retrieved_docs = retriever.invoke(question)
    context = _format_docs(retrieved_docs)

    # Build prompt, keeping prior chat history inline
    human_input = f"{formatted_history}\nUser: {question}" if formatted_history else question

    messages = PROMPT_TEMPLATE.format_messages(
        context=context,
        input=human_input,
    )

    response = ""
    # Stream response from Groq
    for chunk in llm.stream(messages):
        content = chunk.content
        if content:
            response += content
            print(content, end="", flush=True)
            yield content
    print(flush=True)

    # Save to memory
    memory.save_context({"input": question}, {"output": response})


=======
>>>>>>> af8b08ecffec43e5dc1c5577f5759f7b6403b00a
# 🔹 CLI loop for testing
if __name__ == "__main__":
    print("🤖 Groq Chatbot (Hybrid RAG + AstraDB) is ready! Type 'exit' to quit.")
    while True:
        user_input = input("You: ")
        if user_input.lower() in ["exit", "quit"]:
            break
        answer = getAnswer(user_input)
        print("Bot:", answer)